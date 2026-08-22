# ╔══════════════════════════════════════════════════════════════╗
# ║     論文寫作助手 - 題目強化・文獻不卡關版                      ║
# ║     建議檔名：thesis_external_api_v5.py                        ║
# ╚══════════════════════════════════════════════════════════════╝
#
# 安裝：
#   pip install streamlit requests pandas python-docx
#
# 執行：
#   streamlit run thesis_external_api.py
#
# API Key 不再寫死在程式碼內。
# 可在側邊欄輸入，或使用環境變數：
#   export GROQ_API_KEY="你的 Groq Key"
#   export GEMINI_API_KEY="你的 Gemini Key"

import os
import io
import json
import re
import hashlib
import time
from typing import Any

import pandas as pd
import requests
import streamlit as st
from docx import Document


# ─────────────────────────────────────────────
# 基本設定
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="論文寫作助手",
    layout="wide",
    page_icon="🎓",
)

GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODELS_URL = "https://api.groq.com/openai/v1/models"
GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta/models"

DEFAULT_GROQ_MODEL = "openai/gpt-oss-120b"
DEFAULT_GEMINI_MODEL = "gemini-2.5-flash"

REQUEST_TIMEOUT = 90


# ─────────────────────────────────────────────
# Session 初始化
# ─────────────────────────────────────────────
defaults = {
    "step": 0,
    "final_title": "",
    "research_topic": "",
    "refs_list": [],
    "sim_data": None,
    "content": {},
    "abstract": "",
    "ai_cache": {},
    "ai_mode": "Groq（免費雲端）",
    "groq_key": os.getenv("GROQ_API_KEY", ""),
    "gemini_key": os.getenv("GEMINI_API_KEY", ""),
    "groq_model": DEFAULT_GROQ_MODEL,
    "gemini_model": DEFAULT_GEMINI_MODEL,
    "research_method": "問卷調查",
    "rate_limit_info": {},
    "economy_mode": True,
    "groq_available_models": [],
    "groq_models_error": "",
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ─────────────────────────────────────────────
# 工具函數
# ─────────────────────────────────────────────
def clean_api_key(value: str) -> str:
    """去除使用者不小心貼入的空白、換行與引號。"""
    if not value:
        return ""
    return value.strip().strip('"').strip("'")


def make_cache_key(
    prompt: str,
    sys_role: str,
    ai_mode: str,
    model_name: str,
    max_tokens: int,
) -> str:
    """使用穩定雜湊，避免 Python hash() 跨執行不一致與碰撞問題。"""
    raw = json.dumps(
        {
            "prompt": prompt,
            "sys_role": sys_role,
            "ai_mode": ai_mode,
            "model": model_name,
            "max_tokens": max_tokens,
        },
        ensure_ascii=False,
        sort_keys=True,
    )
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def extract_json(text: str) -> Any:
    """容錯解析 AI 回傳的 JSON，支援 ```json ... ``` 包裹。"""
    if not text:
        raise ValueError("AI 未回傳內容")

    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        # 嘗試擷取第一個 JSON object / array
        obj_match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        arr_match = re.search(r"\[.*\]", cleaned, flags=re.DOTALL)

        candidates = []
        if obj_match:
            candidates.append(obj_match.group(0))
        if arr_match:
            candidates.append(arr_match.group(0))

        for candidate in candidates:
            try:
                return json.loads(candidate)
            except json.JSONDecodeError:
                continue

    raise ValueError("AI 回傳內容不是有效 JSON")


def parse_error_response(res: requests.Response) -> str:
    """盡量顯示 API 真正的錯誤訊息，但不暴露 API Key。"""
    try:
        data = res.json()
        if isinstance(data, dict):
            err = data.get("error")
            if isinstance(err, dict):
                return str(err.get("message") or err)[:300]
            if err:
                return str(err)[:300]
    except Exception:
        pass

    text = (res.text or "").strip()
    return text[:300] if text else f"HTTP {res.status_code}"


def fetch_groq_models(api_key: str):
    """
    使用目前輸入的 Groq API Key 取得「該帳號實際可用」模型。
    這比把模型 ID 寫死在程式碼安全，能避免模型下架後持續 404。
    """
    api_key = clean_api_key(api_key)
    if not api_key:
        return [], "尚未輸入 Groq API Key"

    try:
        res = requests.get(
            GROQ_MODELS_URL,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            timeout=20,
        )

        if res.status_code == 200:
            data = res.json().get("data", [])
            model_ids = sorted(
                {
                    m.get("id", "").strip()
                    for m in data
                    if isinstance(m, dict) and m.get("id")
                }
            )
            return model_ids, ""

        if res.status_code in (401, 403):
            return [], "Groq API Key 無效或無權限"

        return [], f"Groq 模型清單讀取失敗：{res.status_code} {parse_error_response(res)}"

    except requests.RequestException as e:
        return [], f"Groq 模型清單連線失敗：{str(e)[:150]}"


def choose_recommended_groq_model(model_ids):
    """
    優先選擇適合長文學術寫作的模型。
    Groq 2026-08-16 已停止 llama-3.3-70b-versatile。
    """
    preferred = [
        "openai/gpt-oss-120b",
        "qwen/qwen3.6-27b",
        "openai/gpt-oss-20b",
    ]
    for model_id in preferred:
        if model_id in model_ids:
            return model_id
    return model_ids[0] if model_ids else DEFAULT_GROQ_MODEL


# ─────────────────────────────────────────────
# AI 呼叫
# ─────────────────────────────────────────────
def call_ai(
    prompt: str,
    sys_role: str = "你是嚴謹的學術研究助手。請使用繁體中文回答；不得捏造資料、統計數字或不存在的文獻。",
    max_tokens: int = 4000,
) -> str:
    ai_mode = st.session_state.ai_mode

    if ai_mode == "Groq（免費雲端）":
        api_key = clean_api_key(st.session_state.groq_key)
        model_name = st.session_state.groq_model.strip() or DEFAULT_GROQ_MODEL
    else:
        api_key = clean_api_key(st.session_state.gemini_key)
        model_name = st.session_state.gemini_model.strip() or DEFAULT_GEMINI_MODEL

    if not api_key:
        return f"❌ 尚未輸入 {ai_mode.split('（')[0]} API Key，請到左側「API Key」欄位輸入。"

    cache_key = make_cache_key(prompt, sys_role, ai_mode, model_name, max_tokens)
    if cache_key in st.session_state.ai_cache:
        return st.session_state.ai_cache[cache_key]

    try:
        if ai_mode == "Groq（免費雲端）":
            def do_groq_request():
                response = requests.post(
                    GROQ_API_URL,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": model_name,
                        "messages": [
                            {"role": "system", "content": sys_role},
                            {"role": "user", "content": prompt},
                        ],
                        "max_completion_tokens": max_tokens,
                        "temperature": 0.5,
                    },
                    timeout=REQUEST_TIMEOUT,
                )

                # Groq 會在 response headers 告知目前限額狀態
                st.session_state.rate_limit_info = {
                    "remaining_requests": response.headers.get(
                        "x-ratelimit-remaining-requests", ""
                    ),
                    "remaining_tokens": response.headers.get(
                        "x-ratelimit-remaining-tokens", ""
                    ),
                    "reset_requests": response.headers.get(
                        "x-ratelimit-reset-requests", ""
                    ),
                    "reset_tokens": response.headers.get(
                        "x-ratelimit-reset-tokens", ""
                    ),
                    "retry_after": response.headers.get("retry-after", ""),
                }
                return response

            res = do_groq_request()

            # 若只是短時間 TPM/RPM 限制，自動等候一次再重試。
            if res.status_code == 429:
                retry_after_raw = res.headers.get("retry-after", "")
                try:
                    retry_after = float(retry_after_raw)
                except (TypeError, ValueError):
                    retry_after = 0

                if 0 < retry_after <= 20:
                    time.sleep(retry_after + 0.5)
                    res = do_groq_request()

            if res.status_code == 200:
                data = res.json()
                result = data["choices"][0]["message"]["content"].strip()
            elif res.status_code == 401:
                return "❌ Groq API Key 無效或未授權，請重新確認。"
            elif res.status_code == 403:
                return "❌ Groq 拒絕此模型或專案權限，請確認模型權限設定。"
            elif res.status_code == 404:
                return (
                    f"❌ Groq 找不到模型「{model_name}」或你的帳號沒有存取權。"
                    "請在左側按「🔄 重新讀取 Groq 模型」，再選擇目前可用模型。"
                )
            elif res.status_code == 429:
                retry = res.headers.get("retry-after", "")
                detail = parse_error_response(res)
                retry_text = f"；約 {retry} 秒後可再試" if retry else ""
                return (
                    "⚠️ Groq 目前撞到速率限制（不一定是今日額度用完）"
                    f"{retry_text}。\n\n詳細訊息：{detail}"
                )
            else:
                return f"❌ Groq 錯誤 {res.status_code}：{parse_error_response(res)}"

        else:
            url = f"{GEMINI_API_BASE}/{model_name}:generateContent"
            res = requests.post(
                url,
                headers={
                    "x-goog-api-key": api_key,
                    "Content-Type": "application/json",
                },
                json={
                    "system_instruction": {
                        "parts": [{"text": sys_role}]
                    },
                    "contents": [
                        {
                            "role": "user",
                            "parts": [{"text": prompt}],
                        }
                    ],
                    "generationConfig": {
                        "temperature": 0.5,
                        "maxOutputTokens": max_tokens,
                    },
                },
                timeout=REQUEST_TIMEOUT,
            )

            if res.status_code == 200:
                data = res.json()
                candidates = data.get("candidates", [])
                if not candidates:
                    return "❌ Gemini 未回傳可用內容。"
                parts = candidates[0].get("content", {}).get("parts", [])
                result = "".join(
                    str(p.get("text", "")) for p in parts if isinstance(p, dict)
                ).strip()
                if not result:
                    return "❌ Gemini 回傳內容為空。"
            elif res.status_code == 400:
                return f"❌ Gemini 請求格式或模型錯誤：{parse_error_response(res)}"
            elif res.status_code in (401, 403):
                return "❌ Gemini API Key 無效、未授權，或此專案無模型使用權限。"
            elif res.status_code == 429:
                return "⚠️ Gemini 已達速率或額度限制，請稍後再試或切換 Groq。"
            else:
                return f"❌ Gemini 錯誤 {res.status_code}：{parse_error_response(res)}"

        st.session_state.ai_cache[cache_key] = result
        return result

    except requests.Timeout:
        return "❌ API 連線逾時，請稍後再試。"
    except requests.RequestException as e:
        return f"❌ API 連線錯誤：{str(e)[:200]}"
    except (KeyError, IndexError, TypeError, ValueError) as e:
        return f"❌ API 回傳格式異常：{str(e)[:200]}"


def test_current_api() -> str:
    """用極短請求測試目前選取的 API。"""
    return call_ai(
        "只回答：OK",
        sys_role="你是 API 連線測試器。",
        max_tokens=20,
    )


# ─────────────────────────────────────────────
# 文獻搜尋
# ─────────────────────────────────────────────
def fetch_refs_fast(query: str, total: int = 15):
    """
    只收集實際搜尋到的文獻。
    不使用 AI 生成參考文獻，以避免虛構引用。
    """
    try:
        res = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={
                "query": query,
                "limit": total,
                "fields": "title,authors,year,abstract,venue,url,externalIds",
            },
            headers={"User-Agent": "ThesisWritingAssistant/2.0"},
            timeout=15,
        )

        if res.status_code == 200:
            refs = []
            for p in res.json().get("data", []):
                external_ids = p.get("externalIds") or {}
                refs.append(
                    {
                        "title": p.get("title") or "",
                        "authors": ", ".join(
                            [
                                a.get("name", "")
                                for a in (p.get("authors") or [])[:6]
                                if a.get("name")
                            ]
                        ),
                        "year": p.get("year") or "",
                        "venue": p.get("venue") or "",
                        "doi": external_ids.get("DOI") or "",
                        "url": p.get("url") or "",
                        "abstract": (p.get("abstract") or "")[:1000],
                        "source": "Semantic Scholar",
                    }
                )
            return refs

        if res.status_code == 429:
            st.warning("Semantic Scholar 搜尋頻率過高，請稍後再試。")
        else:
            st.warning(f"文獻搜尋失敗：HTTP {res.status_code}")

    except requests.RequestException as e:
        st.warning(f"文獻搜尋連線失敗：{str(e)[:120]}")

    return []


def build_literature_queries(title: str, topic: str = ""):
    """把中文研究題目轉成適合 Semantic Scholar 的精簡英文檢索詞。"""
    source_text = topic.strip() or title.strip()
    prompt = f"""
研究主題：{source_text}
論文題目：{title}

請整理成 3 組適合 Semantic Scholar 的英文學術搜尋式。
要求：
- 每組 3～6 個核心英文關鍵詞
- 不要逐字翻譯整個中文標題
- 涵蓋主要變項、研究對象或情境
- 只輸出合法 JSON：
{{"queries":["query 1","query 2","query 3"]}}
"""
    result = call_ai(
        prompt,
        sys_role="你是學術資料庫檢索專家。只輸出合法 JSON。",
        max_tokens=220,
    )
    if result.startswith(("❌", "⚠️")):
        return []
    try:
        data = extract_json(result)
        queries = data.get("queries", []) if isinstance(data, dict) else []
        return [q.strip() for q in queries if isinstance(q, str) and q.strip()][:3]
    except Exception:
        return []


def fetch_refs_multi(title: str, topic: str = "", total_target: int = 15):
    """用多組核心關鍵字搜尋真實文獻並去重。"""
    search_queries = build_literature_queries(title, topic)
    fallback_query = topic.strip() or title.strip()
    if fallback_query:
        search_queries.append(fallback_query)
    if not search_queries:
        search_queries = [title]

    refs = []
    seen = set()
    for q in search_queries[:4]:
        results = fetch_refs_fast(q, total=8)
        for ref in results:
            key = ((ref.get("doi") or "").lower().strip()
                   or re.sub(r"\W+", "", (ref.get("title") or "").lower()))
            if not key or key in seen:
                continue
            seen.add(key)
            refs.append(ref)
            if len(refs) >= total_target:
                return refs
    return refs


# ─────────────────────────────────────────────
# 研究模型
# ─────────────────────────────────────────────
def simulate_model_simple(topic: str, method: str):
    prompt = f"""
研究題目：{topic}
研究方法：{method}

請提出一個「概念性研究模型草稿」，只能作為研究設計建議，不能假裝是真實分析結果。

請只輸出有效 JSON，格式如下：
{{
  "criteria": ["構面1", "構面2", "構面3"],
  "weights": [0.34, 0.33, 0.33],
  "note": "權重僅為示意，正式研究必須由實際資料或指定方法計算"
}}
"""
    result = call_ai(
        prompt,
        sys_role="你是研究方法專家。只輸出合法 JSON，不得虛構實證結果。",
        max_tokens=800,
    )

    if result.startswith(("❌", "⚠️")):
        return {
            "criteria": [],
            "weights": [],
            "note": result,
        }

    try:
        data = extract_json(result)
        if isinstance(data, dict):
            data.setdefault(
                "note",
                "此模型僅為概念草稿，正式權重須由真實研究資料計算。",
            )
            return data
    except ValueError:
        pass

    return {
        "criteria": [],
        "weights": [],
        "note": "AI 未回傳有效 JSON，請重新生成。",
    }


# ─────────────────────────────────────────────
# 章節 Prompt
# ─────────────────────────────────────────────
CHAPTER_PROMPTS = {
    "第一章 緒論": """
撰寫第一章緒論草稿，包含：
1. 研究背景與動機
2. 研究目的
3. 研究問題
4. 研究範圍與限制
不得捏造統計數據。若需要特定數據，請標記【待補實際資料】。
""",
    "第二章 文獻探討": """
撰寫第二章文獻探討草稿，包含：
1. 核心理論
2. 主要變項與概念
3. 相關實證研究整理
4. 研究缺口
5. 可驗證之研究假設或研究命題
只能引用下方提供的真實文獻；資料不足時請明確寫【需補充文獻】。
""",
    "第三章 研究方法": """
撰寫第三章研究方法草稿，包含：
1. 研究架構
2. 研究設計
3. 研究對象與抽樣
4. 變項操作化
5. 資料蒐集
6. 分析方法
未知的樣本數、量表來源、信效度門檻或分析結果不得自行填造。
""",
    "第四章 結果": """
只建立第四章「結果章撰寫架構與範本」，不得虛構任何樣本數、平均數、
標準差、p 值、迴歸係數、模型適配度、權重或其他分析結果。
凡需要真實研究資料的位置，一律以【待填實際分析結果】標示。
""",
    "第五章 結論": """
建立第五章結論與建議的草稿框架，包含：
1. 研究發現摘要
2. 理論貢獻
3. 實務貢獻
4. 研究限制
5. 後續研究建議
若第四章尚無真實數據，不得聲稱任何假設獲得支持。
""",
}


def make_refs_text(refs, limit=5, include_abstract=False):
    """
    省額度版文獻上下文：
    - 一般章節只傳作者、年份、題名、期刊
    - 只有第二章文獻探討才附短摘要
    - 避免每章重複塞入大量 abstract，造成 TPM 快速耗盡
    """
    lines = []
    for i, r in enumerate(refs[:limit], start=1):
        title = r.get("title", "")
        authors = r.get("authors", "")
        year = r.get("year", "")
        venue = r.get("venue", "")
        doi = r.get("doi", "")

        line = f"[{i}] {authors} ({year}). {title}. {venue}."
        if doi:
            line += f" DOI:{doi}"

        if include_abstract:
            abstract = (r.get("abstract") or "").strip()
            if abstract:
                # 只保留短摘要，避免輸入 token 暴增
                line += f"\n摘要片段：{abstract[:300]}"

        lines.append(line)

    return "\n\n".join(lines)

def write_chapter_fast(chapter: str, title: str, refs):
    # 只有文獻探討需要摘要；其他章節只傳精簡書目資訊
    include_abstract = chapter == "第二章 文獻探討"
    refs_text = make_refs_text(
        refs,
        limit=5,
        include_abstract=include_abstract,
    )

    if chapter == "第二章 文獻探討" and not refs:
        refs_text = (
            "目前尚無可驗證文獻。請建立完整文獻探討架構，"
            "但所有需要引用之處一律標示【待補文獻】，不得自行杜撰作者或年份。"
        )

    prompt = f"""
論文題目：{title}

可使用的真實文獻：
{refs_text if refs_text else "目前未提供文獻。"}

任務：
{CHAPTER_PROMPTS.get(chapter, "")}

寫作要求：
- 使用正式、自然的繁體中文學術語氣。
- 不得杜撰作者、年份、期刊、DOI、研究數據或研究結論。
- 若資訊不足，以【待補資料】清楚標記。
- 引用時僅使用上方列出的文獻。
- 本次輸出為論文草稿，不宣稱已完成正式研究。
"""

    # 舊版大多是 2000 tokens；這裡依章節控制，
    # 避免五章連續各 3500 tokens 把免費 TPM 撞滿。
    economy_limits = {
        "第一章 緒論": 1400,
        "第二章 文獻探討": 2000,
        "第三章 研究方法": 1600,
        "第四章 結果": 1200,
        "第五章 結論": 1200,
    }
    normal_limits = {
        "第一章 緒論": 1800,
        "第二章 文獻探討": 2400,
        "第三章 研究方法": 2000,
        "第四章 結果": 1500,
        "第五章 結論": 1500,
    }

    limits = economy_limits if st.session_state.economy_mode else normal_limits
    return call_ai(prompt, max_tokens=limits.get(chapter, 1600))


def generate_abstract():
    """根據已產生的章節生成摘要；沒有真實結果時要明確保留。"""
    chapter_text = "\n\n".join(
        f"{ch}\n{st.session_state.content.get(ch, '')}"
        for ch in CHAPTER_PROMPTS
    )

    prompt = f"""
論文題目：{st.session_state.final_title}

以下是目前論文草稿：
{chapter_text[:10000]}

請撰寫 400～600 字繁體中文摘要草稿。
必須包含研究目的、方法、預期或已知結果狀態、研究價值。
若第四章沒有真實數據，請用「本研究結果尚待實證資料分析完成」之類的表述，
不得虛構研究發現。
"""
    return call_ai(prompt, max_tokens=800)


def format_reference(ref: dict) -> str:
    authors = ref.get("authors", "").strip() or "作者資料未提供"
    year = ref.get("year", "") or "n.d."
    title = ref.get("title", "").strip()
    venue = ref.get("venue", "").strip()
    doi = ref.get("doi", "").strip()

    text = f"{authors} ({year}). {title}."
    if venue:
        text += f" {venue}."
    if doi:
        text += f" https://doi.org/{doi}"
    return text


# ─────────────────────────────────────────────
# 主介面
# ─────────────────────────────────────────────
st.title("🎓 論文寫作助手")
st.caption("🔐 API Key 由使用者於外部輸入，不儲存在程式碼內")

with st.sidebar:
    st.header("⚙️ AI 設定")

    st.session_state.ai_mode = st.radio(
        "選擇 AI 服務",
        ["Groq（免費雲端）", "Gemini"],
        key="ai_mode_selector",
        help="切換後，請在下方輸入對應 API Key。",
    )

    st.subheader("🔑 API Key")

    st.text_input(
        "Groq API Key",
        type="password",
        key="groq_key",
        placeholder="gsk_...",
        help="只保存在目前 Streamlit 工作階段；也可使用 GROQ_API_KEY 環境變數。",
    )

    st.text_input(
        "Gemini API Key",
        type="password",
        key="gemini_key",
        placeholder="AIza...",
        help="只保存在目前 Streamlit 工作階段；也可使用 GEMINI_API_KEY 環境變數。",
    )

    with st.expander("🧠 模型設定", expanded=True):
        if st.session_state.ai_mode == "Groq（免費雲端）":
            if st.button("🔄 重新讀取 Groq 模型"):
                models, err = fetch_groq_models(st.session_state.groq_key)
                st.session_state.groq_available_models = models
                st.session_state.groq_models_error = err

                if models:
                    recommended = choose_recommended_groq_model(models)
                    st.session_state.groq_model = recommended
                    st.success(f"已取得 {len(models)} 個可用模型")
                elif err:
                    st.error(err)

            # 第一次輸入 Key 後若尚未抓過清單，自動嘗試一次
            if (
                clean_api_key(st.session_state.groq_key)
                and not st.session_state.groq_available_models
                and not st.session_state.groq_models_error
            ):
                models, err = fetch_groq_models(st.session_state.groq_key)
                st.session_state.groq_available_models = models
                st.session_state.groq_models_error = err

                if models:
                    recommended = choose_recommended_groq_model(models)
                    if st.session_state.groq_model not in models:
                        st.session_state.groq_model = recommended

            available = st.session_state.groq_available_models

            if available:
                # 確保目前值存在於 selectbox options
                if st.session_state.groq_model not in available:
                    st.session_state.groq_model = choose_recommended_groq_model(available)

                st.selectbox(
                    "Groq 模型",
                    options=available,
                    key="groq_model",
                    help="此清單由 Groq API 即時讀取，顯示你的帳號目前實際可用模型。",
                )

                recommended = choose_recommended_groq_model(available)
                st.caption(f"📌 長文寫作建議：{recommended}")
            else:
                st.text_input(
                    "Groq 模型 ID",
                    key="groq_model",
                    help=f"目前無法取得模型清單；預設：{DEFAULT_GROQ_MODEL}",
                )
                if st.session_state.groq_models_error:
                    st.caption(st.session_state.groq_models_error)

        else:
            st.text_input(
                "Gemini 模型 ID",
                key="gemini_model",
                help=f"預設：{DEFAULT_GEMINI_MODEL}",
            )

    selected_key = (
        st.session_state.groq_key
        if st.session_state.ai_mode == "Groq（免費雲端）"
        else st.session_state.gemini_key
    )
    st.caption("✅ 已輸入 API Key" if clean_api_key(selected_key) else "⚠️ 尚未輸入目前服務的 API Key")

    if st.button("🧪 測試 API 連線"):
        with st.spinner("測試中..."):
            test_result = test_current_api()
        if test_result.strip().upper() == "OK":
            st.success("API 連線正常")
        elif test_result.startswith("❌"):
            st.error(test_result)
        else:
            st.info(test_result)

    if st.button("🧹 清除 AI 快取"):
        st.session_state.ai_cache = {}
        st.success("已清除快取")

    st.toggle(
        "💰 省額度模式",
        key="economy_mode",
        help="縮短每章輸出並減少送入 API 的文獻內容；建議 Groq 免費方案開啟。",
    )

    if st.session_state.ai_mode == "Groq（免費雲端）":
        info = st.session_state.rate_limit_info
        if info:
            with st.expander("📊 Groq 即時限額狀態"):
                remaining_tokens = info.get("remaining_tokens") or "—"
                remaining_requests = info.get("remaining_requests") or "—"
                reset_tokens = info.get("reset_tokens") or "—"
                st.write(f"本分鐘剩餘 tokens：{remaining_tokens}")
                st.write(f"剩餘 requests：{remaining_requests}")
                st.write(f"Token 重置：約 {reset_tokens}")
                st.caption(
                    "這些數值來自 Groq API 回應；429 不一定代表每日額度用完。"
                )

    st.header("📝 研究設定")
    st.selectbox(
        "研究方法",
        ["問卷調查", "個案研究", "文獻分析", "混合方法"],
        key="research_method",
    )

    st.warning(
        "學術提醒：本工具只協助產生草稿與研究設計。"
        "實證結果、樣本數、統計值與參考文獻都必須由真實資料驗證。"
    )


# ─────────────────────────────────────────────
# 進度指示
# ─────────────────────────────────────────────
steps = ["題目", "文獻", "模型", "寫作"]
cols = st.columns(4)

for i, step in enumerate(steps):
    with cols[i]:
        if st.session_state.step > i:
            st.success(f"✓ {step}")
        elif st.session_state.step == i:
            st.info(f"→ {step}")
        else:
            st.caption(step)


# ─────────────────────────────────────────────
# 步驟 0：題目
# ─────────────────────────────────────────────
if st.session_state.step == 0:
    st.subheader("📝 步驟 1：研究題目")

    topic = st.text_input(
        "輸入研究主題（例如：AI 教育、ESG 投資）",
        placeholder="人才培訓與教育訓練",
    )

    if st.button("💡 生成題目建議"):
        if not topic.strip():
            st.warning("請先輸入研究主題。")
        else:
            with st.spinner("AI 生成中..."):
                method_now = st.session_state.research_method
                result = call_ai(
                    f"""
你是一位熟悉台灣碩士論文審查標準的研究指導教授。

研究主題：{topic}
偏好研究方法：{method_now}

請先判斷這個主題可能的研究缺口，再提出 5 個真正可執行、具有研究價值，而且不要太普通的碩士論文題目。

出題要求：
1. 避免只寫成空泛的「A 對 B 之影響研究」。
2. 題目要有明確研究情境、研究對象或產業範圍。
3. 變項之間必須有合理理論關係；不要為了看起來厲害而硬塞中介或調節。
4. 優先找出實務痛點、理論缺口或新情境，而不是只換變項名稱。
5. 題名要像正式碩士論文，不要像一般報告或企劃案。
6. 要兼顧資料可取得性、問卷或資料來源可行性，以及研究者在碩士階段能完成。
7. 若使用者偏好的研究方法不適合某題，可提出較合適的方法並簡短說明。
8. 五題之間要有明顯差異，不要只是同一句換字。

請依這個格式輸出：

【推薦 1｜整體最值得做】
題目：
研究亮點：
研究對象：
建議方法：
為什麼比一般題目好：

【推薦 2｜實務性最高】
...

【推薦 3｜創新性最高】
...

【推薦 4｜最容易執行】
...

【推薦 5｜備選】
...

最後加一行：
如果要投稿期刊，我最推薦：第 X 題，原因：……
"""
                )
            st.write(result)

    title = st.text_input(
        "✅ 確認題目",
        st.session_state.final_title,
        placeholder="例如：生成式 AI 使用對員工工作重塑之影響：科技自我效能的調節效果",
    )

    if st.button("🚀 開始建立論文草稿", type="primary"):
        if not title.strip():
            st.warning("請先確認論文題目。")
        else:
            # 若題目改變，清除舊研究內容，避免混用
            if title.strip() != st.session_state.final_title:
                st.session_state.refs_list = []
                st.session_state.sim_data = None
                st.session_state.content = {}
                st.session_state.abstract = ""

            st.session_state.research_topic = topic.strip()
            st.session_state.final_title = title.strip()
            st.session_state.step = 1
            st.rerun()


# ─────────────────────────────────────────────
# 步驟 1：文獻
# ─────────────────────────────────────────────
elif st.session_state.step == 1:
    st.subheader("📚 步驟 2：文獻收集")
    st.info(f"📌 研究題目：{st.session_state.final_title}")

    st.caption(
        "此版本只收錄資料庫實際搜尋到的文獻；AI 只協助產生檢索關鍵字，不會生成假文獻。"
    )

    if st.button("🔍 搜尋真實文獻", type="primary"):
        with st.spinner("搜尋中..."):
            refs = fetch_refs_multi(
                st.session_state.final_title,
                st.session_state.research_topic,
                total_target=15,
            )
            st.session_state.refs_list = refs

        if refs:
            st.success(f"✅ 收集到 {len(refs)} 篇可追溯文獻")
        else:
            st.warning("目前沒有取得文獻，但仍可繼續下一步；第二章會以【待補文獻】標示。")

    if st.session_state.refs_list:
        df = pd.DataFrame(st.session_state.refs_list)

        display_cols = [
            c for c in ["title", "authors", "year", "venue", "doi", "source"]
            if c in df.columns
        ]
        st.dataframe(df[display_cols], use_container_width=True)

        st.caption(
            "提醒：搜尋結果仍需由研究者逐篇核對原文、研究設計與引用格式。"
        )

    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ 上一步"):
            st.session_state.step = 0
            st.rerun()

    with col2:
        if st.button("➡️ 下一步", type="primary"):
            if not st.session_state.refs_list:
                st.info(
                    "目前尚未找到文獻，仍可繼續。後續文獻探討會標示【待補文獻】，"
                    "不會自動捏造引用。"
                )
            st.session_state.step = 2
            st.rerun()


# ─────────────────────────────────────────────
# 步驟 2：研究模型
# ─────────────────────────────────────────────
elif st.session_state.step == 2:
    st.subheader("🔬 步驟 3：建立研究模型")
    st.info(f"📌 研究題目：{st.session_state.final_title}")

    if st.button("🛠️ 生成概念性研究模型", type="primary"):
        with st.spinner("模型生成中..."):
            st.session_state.sim_data = simulate_model_simple(
                st.session_state.final_title,
                st.session_state.research_method,
            )

    if st.session_state.sim_data:
        st.json(st.session_state.sim_data)
        st.caption("此處為研究設計草稿，不代表真實統計分析結果。")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("⬅️ 上一步"):
            st.session_state.step = 1
            st.rerun()

    with col2:
        if st.button("➡️ 下一步", type="primary"):
            st.session_state.step = 3
            st.rerun()


# ─────────────────────────────────────────────
# 步驟 3：寫作
# ─────────────────────────────────────────────
elif st.session_state.step == 3:
    st.subheader("✍️ 步驟 4：撰寫論文草稿")
    st.warning(
        "第四章不會自動捏造研究結果；沒有真實分析資料時，"
        "只會產生章節架構並標示【待填實際分析結果】。"
    )

    chapters = list(CHAPTER_PROMPTS.keys())
    progress = sum(1 for c in chapters if c in st.session_state.content)
    st.progress(progress / len(chapters), text=f"已完成 {progress}/{len(chapters)} 章")

    if st.button("🚀 一鍵生成五章草稿", type="primary"):
        if not st.session_state.refs_list:
            st.info(
                "目前沒有可驗證文獻，仍會繼續生成；"
                "第二章引用位置將標示【待補文獻】。"
            )

        generated_count = 0
        for ch in chapters:
            if ch not in st.session_state.content:
                if (
                    generated_count > 0
                    and st.session_state.ai_mode == "Groq（免費雲端）"
                ):
                    time.sleep(3)

                with st.spinner(f"正在撰寫 {ch}..."):
                    result = write_chapter_fast(
                        ch,
                        st.session_state.final_title,
                        st.session_state.refs_list,
                    )
                    st.session_state.content[ch] = result
                    generated_count += 1

                if result.startswith("⚠️ Groq 目前撞到速率限制"):
                    st.warning(
                        "已暫停後續章節，避免持續消耗請求。"
                        "稍後再按一次即可從未完成章節續寫。"
                    )
                    break

        st.success("✅ 本輪可完成的章節已保存")

    for ch in chapters:
        with st.expander(
            f"{'✅' if ch in st.session_state.content else '📝'} {ch}"
        ):
            if ch in st.session_state.content:
                text = st.session_state.content[ch]
                word_count = len(re.sub(r"\s+", "", text))
                st.caption(f"約 {word_count} 字")
                st.write(text)

                if st.button(f"🔄 重新撰寫 {ch}", key=f"rewrite_{ch}"):
                    del st.session_state.content[ch]
                    st.rerun()
            else:
                if st.button(f"✏️ 撰寫 {ch}", key=f"write_{ch}"):
                    with st.spinner(f"正在撰寫 {ch}..."):
                        st.session_state.content[ch] = write_chapter_fast(
                            ch,
                            st.session_state.final_title,
                            st.session_state.refs_list,
                        )
                    st.rerun()

    if all(ch in st.session_state.content for ch in chapters):
        st.divider()
        st.subheader("🧾 摘要")

        if not st.session_state.abstract:
            if st.button("✨ 生成摘要草稿"):
                with st.spinner("生成摘要中..."):
                    st.session_state.abstract = generate_abstract()
                st.rerun()
        else:
            st.write(st.session_state.abstract)
            if st.button("🔄 重新生成摘要"):
                st.session_state.abstract = generate_abstract()
                st.rerun()

        st.divider()
        st.subheader("📥 下載完整論文草稿")

        full_text = f"# {st.session_state.final_title}\n\n"

        full_text += "## 摘要\n\n"
        full_text += (
            st.session_state.abstract
            if st.session_state.abstract
            else "【摘要尚未生成】"
        )
        full_text += "\n\n"

        for ch in chapters:
            full_text += (
                f"## {ch}\n\n"
                f"{st.session_state.content.get(ch, '')}\n\n"
            )

        full_text += "## 參考文獻\n\n"
        for ref in st.session_state.refs_list:
            full_text += f"- {format_reference(ref)}\n"

        total_words = len(re.sub(r"\s+", "", full_text))
        st.metric("📊 草稿總字數", f"{total_words} 字")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "📄 下載 Markdown",
                full_text,
                file_name=f"{st.session_state.final_title[:30]}.md",
                mime="text/markdown",
            )

        with col2:
            txt_text = re.sub(r"^#+\s*", "", full_text, flags=re.MULTILINE)
            st.download_button(
                "📝 下載 TXT",
                txt_text,
                file_name=f"{st.session_state.final_title[:30]}.txt",
                mime="text/plain",
            )

        with col3:
            doc = Document()
            doc.add_heading(st.session_state.final_title, 0)

            doc.add_heading("摘要", level=1)
            doc.add_paragraph(
                st.session_state.abstract
                if st.session_state.abstract
                else "【摘要尚未生成】"
            )

            for ch in chapters:
                doc.add_heading(ch, level=1)
                doc.add_paragraph(st.session_state.content.get(ch, ""))

            doc.add_heading("參考文獻", level=1)
            for ref in st.session_state.refs_list:
                doc.add_paragraph(
                    format_reference(ref),
                    style="List Bullet",
                )

            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.download_button(
                "📘 下載 Word",
                doc_io,
                file_name=f"{st.session_state.final_title[:30]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    if st.button("⬅️ 返回上一步"):
        st.session_state.step = 2
        st.rerun()


# ─────────────────────────────────────────────
# 頁腳
# ─────────────────────────────────────────────
st.divider()
st.caption(
    "🔐 API Key 不寫入原始碼。建議正式部署時使用環境變數或平台 Secrets 管理。"
)
